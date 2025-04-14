import PyPDF2
import docx
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import os
import uuid
import logging

logger = logging.getLogger(__name__)

#todo add more logging to functions
def extract_text(file_path):
    """Extract text from PDF or DOCX file."""
    logger.info(f"Extracting text from {file_path}")
    try:
        if file_path.endswith(".pdf"):
            return extract_pdf_text(file_path)
        elif file_path.endswith(".docx"):
            return extract_docx_text(file_path)
        logger.warning(f"Unsupported file format for {file_path}")
        return ""
    except Exception as e:
        logger.exception(f"Error extracting text from {file_path}: {str(e)}")
        raise

def extract_pdf_text(file_path):
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        text = "".join([page.extract_text() for page in reader.pages])
    return text

def extract_docx_text(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def create_document(text, output_format="pdf"):
    """
    Create a PDF or DOCX document from text.
    Returns the path to the created document.
    """
    # Create a unique filename
    filename = f"improved_resume_{uuid.uuid4()}"
    output_dir = "media/improved_resumes"
    
    # Ensure directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    if output_format.lower() == "pdf":
        output_path = os.path.join(output_dir, f"{filename}.pdf")
        create_pdf(text, output_path)
    else:
        output_path = os.path.join(output_dir, f"{filename}.docx")
        create_docx(text, output_path)
        
    return output_path

def create_pdf(text, output_path):
    """Create a PDF file from text."""
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    
    # Set font and size
    c.setFont("Helvetica", 10)
    
    # Split text into lines for easier processing
    lines = text.split("\n")
    
    # Variables to track position
    y_position = height - 50
    line_height = 12
    
    for line in lines:
        # Check if we need a new page
        if y_position < 50:
            c.showPage()
            y_position = height - 50
            c.setFont("Helvetica", 10)
        
        # Write the line and move down
        c.drawString(50, y_position, line)
        y_position -= line_height
    
    c.save()

def create_docx(text, output_path):
    """Create a DOCX file from text."""
    doc = docx.Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process text by paragraphs
    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph()
            lines = para_text.split("\n")
            p.add_run(lines[0])
            
            # Add remaining lines as separate runs with line breaks
            for line in lines[1:]:
                if line.strip():
                    p.add_run("\n" + line)
    
    doc.save(output_path)